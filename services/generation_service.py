import os
import re
import requests
from typing import Dict, List, Optional
from services.retrieval_service import RetrievalService
from services.persistence_service import PersistenceService


class GenerationService:
    DOCUMENT_TYPES = {
        "informe": {
            "label": "Informe Técnico",
            "icon": "bi-clipboard-data",
            "system": """Eres un redactor experto en documentos institucionales universitarios peruanos.
Generas informes técnicos formales con estructura clara y lenguaje profesional.
SIEMPRE usa este esquema:
# [TÍTULO DEL INFORME]
## I. INTRODUCCIÓN
## II. ANTECEDENTES
## III. ANÁLISIS / DESARROLLO
## IV. CONCLUSIONES
## V. RECOMENDACIONES
Usa numeración romana para secciones. Lenguaje formal. Sin markdown decorativo innecesario.""",
            "instruction": "Genera un informe técnico institucional completo y bien estructurado basado en las instrucciones."
        },
        "acta": {
            "label": "Acta de Reunión",
            "icon": "bi-journal-text",
            "system": """Eres un secretario institucional experto en redacción de actas formales universitarias peruanas.
SIEMPRE usa este esquema:
# ACTA N° [NÚMERO] - [NOMBRE DEL ÓRGANO]
**Fecha:** [fecha]
**Hora de inicio:** [hora]
**Lugar:** [lugar]
**Presidente:** [nombre]
**Secretario:** [nombre]
## ASISTENTES
## ORDEN DEL DÍA
## DESARROLLO DE LA SESIÓN
### Punto 1: [tema]
## ACUERDOS
## CIERRE
**Hora de cierre:** [hora]
_________________________________
Firmas""",
            "instruction": "Genera un acta de reunión formal y completa basada en las instrucciones."
        },
        "memo": {
            "label": "Memorando",
            "icon": "bi-envelope-paper",
            "system": """Eres un redactor experto en comunicaciones internas institucionales universitarias.
SIEMPRE usa este esquema exacto:
MEMORANDO N° [NÚMERO]-[AÑO]-[UNIDAD]
**PARA:** [destinatario y cargo]
**DE:** [remitente y cargo]  
**ASUNTO:** [asunto concreto]
**FECHA:** [fecha]
---
[Cuerpo del memorando: saludo protocolar, desarrollo del asunto, petición o comunicado]
Atentamente,
[Firma y cargo]""",
            "instruction": "Genera un memorando institucional formal basado en las instrucciones."
        },
        "resolucion": {
            "label": "Resolución",
            "icon": "bi-bank",
            "system": """Eres un redactor jurídico-administrativo experto en resoluciones universitarias peruanas.
SIEMPRE usa este esquema:
# RESOLUCIÓN [RECTORAL/DECANAL/DIRECTORAL] N° [NÚMERO]-[AÑO]-[SIGLAS]
**[Ciudad], [fecha completa]**
## VISTO:
[Expediente o documento que origina la resolución]
## CONSIDERANDO:
Que, [primer considerando];
Que, [segundo considerando];
Que, [estando a lo expuesto y en uso de las atribuciones conferidas];
## RESUELVE:
**Artículo 1°.-** [Primera disposición]
**Artículo 2°.-** [Segunda disposición]
**Artículo 3°.-** Regístrese, comuníquese y archívese.
[Firma y cargo]""",
            "instruction": "Genera una resolución administrativa formal basada en las instrucciones."
        },
        "oficio": {
            "label": "Oficio",
            "icon": "bi-send",
            "system": """Eres un redactor experto en comunicaciones oficiales institucionales universitarias peruanas.
SIEMPRE usa este esquema:
OFICIO N° [NÚMERO]-[AÑO]-[UNIDAD]
[Ciudad], [fecha]
Señor(a):
[Nombre completo]
[Cargo]
[Institución]
Presente.-
Asunto: [asunto]
---
[Cuerpo: saludo protocolar, motivo, desarrollo, petición o comunicado]
Es propicia la oportunidad para expresarle las muestras de mi especial consideración y estima.
Atentamente,
[Firma]
[Nombre completo]
[Cargo]""",
            "instruction": "Genera un oficio institucional formal basado en las instrucciones."
        },
        "libre": {
            "label": "Documento Libre",
            "icon": "bi-pen",
            "system": """Eres un redactor experto en documentos institucionales.
Generas documentos bien estructurados, profesionales y coherentes.
REGLAS ESTRICTAS:
- Usa la información de referencia como fuente, NO la copies literalmente
- SIEMPRE incluye los datos específicos del titular o sujeto del documento (nombres, DNI, fechas, etc.)
- Sintetiza y redacta con tus propias palabras
- Estructura con encabezados claros en Markdown
- El documento debe ser coherente de principio a fin
- NO incluyas frases como 'DATOS DEL DOCUMENTO' ni metadatos del sistema
FORMATO OBLIGATORIO: Empieza siempre con:
# [Título descriptivo del documento]""",
            "instruction": "Genera el documento solicitado de forma profesional, sintetizando la información de referencia."
        }
    }

    def __init__(self, retrieval_service: RetrievalService, 
                 persistence_service: PersistenceService):
        self.retrieval = retrieval_service
        self.persistence = persistence_service
        self.base_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
        self.chat_model = os.environ.get("OLLAMA_CHAT_MODEL", "qwen2.5:3b")
        self.generated_dir = "data/generated"
        os.makedirs(self.generated_dir, exist_ok=True)

    def generate(self, prompt: str, doc_type: str = "libre",
                 mode: str = "prompt_libre",
                 source_doc_ids: List[int] = None,
                 doc_format: str = "markdown",
                 user_id: int = None) -> Dict:
        """
        Genera un documento nuevo usando IA con prompts especializados por tipo.
        """
        template = self.DOCUMENT_TYPES.get(doc_type, self.DOCUMENT_TYPES["libre"])
        context_text = ""
        used_doc_ids = []

        # Recuperar contexto si aplica
        if mode in ("basado_repositorio", "basado_documento"):
            doc_id_filter = source_doc_ids[0] if (
                mode == "basado_documento" and source_doc_ids
            ) else None

            # Query expandida para mejor retrieval
            retrieval_query = prompt
            if mode == "basado_documento":
                retrieval_query = f"{prompt}. datos personales nombre apellidos DNI documento identidad fecha nacimiento antecedentes trayectoria educativa experiencia laboral vigencia"

            results = self.retrieval.search(
                query=retrieval_query,
                top_k=10,
                document_id=doc_id_filter,
                sql_threshold=0.25,
                use_rerank=True
            )

            # Para documentos específicos: forzar chunks con datos clave via SQL directo
            forced_chunks = []
            if mode == "basado_documento" and doc_id_filter:
                try:
                    conn = self.persistence.db.get_connection()
                    from psycopg2.extras import RealDictCursor
                    with conn.cursor(cursor_factory=RealDictCursor) as cur:
                        cur.execute("""
                            SELECT chunk_text as text, 1.0 as score, document_id
                            FROM document_chunks
                            WHERE document_id = %s
                            AND (
                                chunk_text ILIKE '%nombre%'
                                OR chunk_text ILIKE '%apellido%'
                                OR chunk_text ILIKE '%nacimiento%'
                                OR chunk_text ILIKE '%DNI%'
                                OR chunk_text ILIKE '%N° de documento%'
                                OR chunk_text ILIKE '%titular%'
                            )
                            ORDER BY chunk_index ASC
                            LIMIT 3;
                        """, (doc_id_filter,))
                        forced_chunks = [dict(r) for r in cur.fetchall()]
                    conn.close()
                except Exception as e:
                    print(f"[GEN] Error forzando chunks: {e}")

            if results or forced_chunks:
                valid = [r for r in results if r.get("score", 0) >= 0.30]
                
                # Combinar: forced_chunks primero, luego el retrieval normal (sin duplicados)
                forced_texts = {c["text"] for c in forced_chunks}
                unique_results = [r for r in valid if r.get("text") not in forced_texts]
                
                combined = forced_chunks + unique_results
                context_parts = [r["text"] for r in combined[:8]]
                
                # Limpiar headers de chunking
                clean_parts = []
                for part in context_parts:
                    cleaned = re.sub(r'^##\s+[A-ZÁÉÍÓÚÑ\s]+\n', '', part, flags=re.MULTILINE)
                    cleaned = cleaned.strip()
                    if cleaned:
                        clean_parts.append(cleaned)
                
                context_text = "\n\n---\n\n".join(clean_parts)
                used_doc_ids = list({r.get("document_id") for r in combined if r.get("document_id")})

        # Construir mensajes para api/chat
        system_msg = template["system"]

        if context_text:
            user_msg = f"""{template['instruction']}

CONTEXTO DE REFERENCIA (usa esta información como base, NO la copies directamente):
{context_text}

INSTRUCCIONES DEL USUARIO:
{prompt}

IMPORTANTE: Redacta el documento de forma coherente y estructurada. NO copies el contexto palabra por palabra. Sintetiza y presenta la información de forma profesional según el tipo de documento solicitado."""
        else:
            user_msg = f"""{template['instruction']}

INSTRUCCIONES DEL USUARIO:
{prompt}"""

        # Llamar al LLM via api/chat (no api/generate)
        try:
            response = requests.post(
                f"{self.base_url}/api/chat",
                json={
                    "model": self.chat_model,
                    "messages": [
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": user_msg}
                    ],
                    "stream": False,
                    "options": {
                        "temperature": 0.4,
                        "top_p": 0.8,
                        "num_predict": 2048
                    }
                },
                timeout=180
            )
            response.raise_for_status()
            content = response.json().get("message", {}).get("content", "").strip()
        except requests.exceptions.Timeout:
            return {"success": False, "error": "⏱️ El modelo tardó demasiado. Intenta con instrucciones más cortas."}
        except Exception as e:
            return {"success": False, "error": str(e)}

        if not content:
            return {"success": False, "error": "El modelo no generó contenido."}

        # Extraer título — buscar primero un encabezado H1, luego H2, luego usar el prompt
        title = ""

        # Buscar línea con # al inicio
        for line in content.strip().split("\n"):
            clean = line.strip()
            if clean.startswith("# "):
                title = clean[2:].strip()
                break
            elif clean.startswith("## "):
                title = clean[3:].strip()
                break

        # Si no hay encabezado markdown, usar las primeras palabras del prompt como título
        if not title:
            # Limpiar el prompt para usarlo como título
            title = prompt.strip()[:80].split("\n")[0]
            # Quitar palabras genéricas al inicio
            title = re.sub(r'^(genera|redacta|crea|elabora|haz|armame|dame)\s+', '', title, flags=re.IGNORECASE).strip()
            title = title[:80]

        # Capitalizar primera letra
        if title:
            title = title[0].upper() + title[1:]

        word_count = len(content.split())

        # Persistir en DB
        try:
            conn = self.persistence.db.get_connection()
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO generated_documents_v2
                (user_id, title, prompt, content, format,
                 generation_mode, source_doc_ids, model_used, word_count)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id;
            """, (
                user_id, title, prompt, content, doc_format,
                mode, source_doc_ids or [],
                self.chat_model, word_count
            ))
            gen_id = cur.fetchone()[0]
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"[ERROR] No se pudo persistir: {e}")
            gen_id = None

        return {
            "success": True,
            "id": gen_id,
            "title": title,
            "content": content,
            "word_count": word_count,
            "mode": mode,
            "doc_type": doc_type,
            "used_doc_ids": used_doc_ids
        }

    def get_all(self, user_id: int = None) -> List[Dict]:
        """Lista todos los documentos generados."""
        conn = self.persistence.db.get_connection()
        from psycopg2.extras import RealDictCursor
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                if user_id:
                    cur.execute("""
                        SELECT id, title, format, generation_mode, 
                               word_count, created_at
                        FROM generated_documents_v2
                        WHERE user_id = %s
                        ORDER BY created_at DESC;
                    """, (user_id,))
                else:
                    cur.execute("""
                        SELECT id, title, format, generation_mode,
                               word_count, created_at
                        FROM generated_documents_v2
                        ORDER BY created_at DESC;
                    """)
                return [dict(r) for r in cur.fetchall()]
        finally:
            conn.close()

    def get_by_id(self, gen_id: int) -> Optional[Dict]:
        """Obtiene un documento generado por ID."""
        conn = self.persistence.db.get_connection()
        from psycopg2.extras import RealDictCursor
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute(
                    "SELECT * FROM generated_documents_v2 WHERE id = %s;",
                    (gen_id,)
                )
                row = cur.fetchone()
                return dict(row) if row else None
        finally:
            conn.close()

    def delete(self, gen_id: int) -> bool:
        """Elimina un documento generado."""
        conn = self.persistence.db.get_connection()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "DELETE FROM generated_documents_v2 WHERE id = %s;",
                    (gen_id,)
                )
            conn.commit()
            return True
        except Exception:
            return False
        finally:
            conn.close()

    def export_docx(self, gen_id: int) -> bytes:
        """Exporta documento generado como DOCX en memoria."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        import re

        doc_data = self.get_by_id(gen_id)
        if not doc_data:
            raise ValueError("Documento no encontrado")

        doc = Document()

        # Estilo del documento
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Procesar contenido línea por línea
        for line in doc_data["content"].split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), 
                                  style='List Number')
            else:
                # Limpiar markdown bold/italic
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                doc.add_paragraph(clean)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_pdf(self, gen_id: int) -> bytes:
        from fpdf import FPDF
        import re

        doc_data = self.get_by_id(gen_id)
        if not doc_data:
            raise ValueError("Documento no encontrado")

        pdf = FPDF(format='A4')
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        # A4 = 210mm ancho. Márgenes 20mm cada lado → área útil = 170mm
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        pdf.set_top_margin(20)
        # Ancho efectivo para multi_cell
        effective_width = 170

        def clean(text):
            replacements = {
                '•': '-', '–': '-', '—': '-',
                '\u2019': "'", '\u2018': "'",
                '\u201c': '"', '\u201d': '"',
                '°': 'o', '→': '->', '←': '<-',
                '\u00b7': '-', '\u2022': '-',
                '\u00ba': 'o', '\u00aa': 'a',
            }
            for char, replacement in replacements.items():
                text = text.replace(char, replacement)
            return text.encode('latin-1', errors='replace').decode('latin-1')

        for line in doc_data["content"].split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(4)
                continue

            if line.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(30, 30, 30)
                pdf.multi_cell(effective_width, 10, clean(line[2:]))
                pdf.ln(3)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(50, 50, 50)
                pdf.multi_cell(effective_width, 8, clean(line[3:]))
                pdf.ln(2)
            elif line.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(70, 70, 70)
                pdf.multi_cell(effective_width, 7, clean(line[4:]))
                pdf.ln(1)
            elif re.match(r'^[-*•]\s+', line):
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(40, 40, 40)
                content_line = re.sub(r'^[-*•]\s+', '', line)
                content_line = re.sub(r'\*\*(.*?)\*\*', r'\1', content_line)
                pdf.multi_cell(effective_width, 6, clean(f"  - {content_line}"))
            elif line.startswith("---"):
                pdf.ln(2)
                pdf.set_draw_color(200, 200, 200)
                pdf.line(20, pdf.get_y(), 190, pdf.get_y())
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(40, 40, 40)
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
                pdf.multi_cell(effective_width, 6, clean(clean_line))

        return bytes(pdf.output())
