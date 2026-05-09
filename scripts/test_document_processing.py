import os
from pathlib import Path
import docx
from fpdf import FPDF
import tempfile
import sys

# Ajustar path temporalmente para importar el servicio desde la carpeta raíz del proyecto
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from services.document_service import DocumentService

def create_dummy_txt(path: str):
    with open(path, "w", encoding="utf-8") as f:
        f.write("Este es un archivo de prueba en formato TXT.\nContiene un poco de texto básico.")

def create_dummy_docx(path: str):
    doc = docx.Document()
    doc.add_heading('Documento de Prueba', 0)
    doc.add_paragraph('Este es un archivo de prueba en formato DOCX.')
    doc.add_paragraph('Contiene múltiples párrafos para verificar la extracción.')
    doc.save(path)

def create_dummy_pdf(path: str):
    # Usando FPDF para crear un PDF simple (requiere fpdf instalado para la prueba)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Documento de Prueba PDF", ln=1, align="C")
    pdf.cell(200, 10, txt="Este es un archivo de prueba en formato PDF.", ln=1)
    pdf.output(path)

def run_tests():
    print("--- INICIANDO PRUEBAS DE PROCESAMIENTO DOCUMENTAL ---")
    doc_service = DocumentService()
    
    # Directorio temporal para los archivos originales antes de ser subidos
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        
        txt_path = temp_path / "prueba.txt"
        docx_path = temp_path / "prueba.docx"
        pdf_path = temp_path / "prueba.pdf"
        
        print("1. Creando archivos de prueba...")
        create_dummy_txt(str(txt_path))
        print("  - TXT creado")
        create_dummy_docx(str(docx_path))
        print("  - DOCX creado")
        try:
            from fpdf import FPDF
            create_dummy_pdf(str(pdf_path))
            print("  - PDF creado")
            files_to_test = [txt_path, docx_path, pdf_path]
        except ImportError:
            print("  - Saltando creación de PDF (fpdf2 no está instalado para la prueba).")
            files_to_test = [txt_path, docx_path]

        print("\n2. Procesando archivos a través del DocumentService")
        for file_path in files_to_test:
            print(f"\nProcesando: {file_path.name}")
            try:
                # El servicio se encarga de todo el flujo
                upload_path, processed_path = doc_service.process_and_save(file_path)
                
                print(f"  [OK] Guardado en uploads: {upload_path}")
                print(f"  [OK] Texto guardado en processed: {processed_path}")
                
                # Leemos un extracto del archivo procesado para verificar
                with open(processed_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    print(f"  [INFO] Extracto (primeros 50 caracteres): {content[:50]}...")
            except Exception as e:
                print(f"  [ERROR] Error procesando {file_path.name}: {e}")

    print("\n--- PRUEBAS FINALIZADAS ---")

if __name__ == "__main__":
    run_tests()
